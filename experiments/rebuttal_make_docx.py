"""
Assemble results/rebuttal/analysis/REBUTTAL_NUMBERS.md and the PNG figures
into a single Word document (the interactive-rebuttal attachment).

Markdown tables become Word tables, '##' headings become Heading 2, bullet
lines become bullets, and every figure in analysis/figures/ is appended to
the section it belongs to.  Requires python-docx (pip install python-docx).

Usage:
    python experiments/rebuttal_make_docx.py
    python experiments/rebuttal_make_docx.py --md results/rebuttal/analysis/REBUTTAL_NUMBERS.md --out rebuttal_attachment.docx
"""

import argparse
import re
from pathlib import Path

FIG_FOR_SECTION = {
    "error bars": "bootstrap_ci.png",
    "failure analysis": "failure_residuals.png",
    "in-sample vs out-of-sample": "u_in_vs_u_out.png",
    "DP-trained": "dp_auc_by_tercile.png",
}


def _add_table(doc, rows):
    from docx.shared import Pt
    header = rows[0]
    body = [r for r in rows[1:] if not set("".join(r)) <= set("-: ")]
    t = doc.add_table(rows=1 + len(body), cols=len(header))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
    for i, row in enumerate(body, start=1):
        for j in range(len(header)):
            cell = t.rows[i].cells[j]
            cell.text = row[j] if j < len(row) else ""
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph()


def _split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def build(md_path: Path, fig_dir: Path, out_path: Path):
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("Supplementary numbers for the interactive rebuttal (#1613)", level=1)
    lines = md_path.read_text().splitlines()
    i, table, current = 0, [], ""
    while i < len(lines):
        line = lines[i]
        if line.startswith("|"):
            table.append(_split_row(line))
            i += 1
            continue
        if table:
            _add_table(doc, table)
            table = []
        if line.startswith("# "):
            pass
        elif line.startswith("## "):
            # figure for the previous section
            for key, fig in FIG_FOR_SECTION.items():
                if key.lower() in current.lower() and (fig_dir / fig).exists():
                    doc.add_picture(str(fig_dir / fig), width=Inches(5.5))
                    doc.add_paragraph(f"Figure: {fig}").italic = True
            current = line[3:].strip()
            doc.add_heading(current, level=2)
        elif line.startswith("- "):
            doc.add_paragraph(re.sub(r"\*\*(.+?)\*\*", r"\1", line[2:]), style="List Bullet")
        elif line.strip():
            p = doc.add_paragraph()
            for part in re.split(r"(\*\*.+?\*\*)", line):
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
        i += 1
    if table:
        _add_table(doc, table)
    for key, fig in FIG_FOR_SECTION.items():
        if key.lower() in current.lower() and (fig_dir / fig).exists():
            doc.add_picture(str(fig_dir / fig), width=Inches(5.5))
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"saved {out_path}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--md", default="results/rebuttal/analysis/REBUTTAL_NUMBERS.md")
    ap.add_argument("--fig_dir", default="results/rebuttal/analysis/figures")
    ap.add_argument("--out", default="results/rebuttal/analysis/rebuttal_attachment.docx")
    args = ap.parse_args()
    try:
        import docx  # noqa: F401
    except ImportError:
        raise SystemExit("pip install python-docx")
    build(Path(args.md), Path(args.fig_dir), Path(args.out))


if __name__ == "__main__":
    main()
